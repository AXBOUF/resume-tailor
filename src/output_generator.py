from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import subprocess
import re
import os
import tempfile
from typing import Dict, Any, List
from src.config import OUTPUT_DIR


# ---------------------------------------------------------------------------
# Jake's Resume LaTeX template helpers
# ---------------------------------------------------------------------------

LATEX_PREAMBLE = r"""\documentclass[letterpaper,11pt]{article}

\usepackage{latexsym}
\usepackage[empty]{fullpage}
\usepackage{titlesec}
\usepackage{marvosym}
\usepackage[usenames,dvipsnames]{color}
\usepackage{verbatim}
\usepackage{enumitem}
\usepackage[hidelinks]{hyperref}
\usepackage{fancyhdr}
\usepackage[english]{babel}
\usepackage{tabularx}
\usepackage{fontspec}

\pagestyle{fancy}
\fancyhf{}
\fancyfoot{}
\renewcommand{\headrulewidth}{0pt}
\renewcommand{\footrulewidth}{0pt}

% Adjust margins
\addtolength{\oddsidemargin}{-0.5in}
\addtolength{\evensidemargin}{-0.5in}
\addtolength{\textwidth}{1in}
\addtolength{\topmargin}{-.5in}
\addtolength{\textheight}{1.0in}

\urlstyle{same}
\raggedbottom
\raggedright
\setlength{\tabcolsep}{0in}

% Sections formatting
\titleformat{\section}{
  \vspace{-4pt}\scshape\raggedright\large
}{}{0em}{}[\color{black}\titlerule \vspace{-5pt}]

%-------------------------
% Custom commands
\newcommand{\resumeItem}[1]{
  \item\small{#1 \vspace{-2pt}}
}

\newcommand{\resumeSubheading}[4]{
  \vspace{-2pt}\item
    \begin{tabular*}{0.97\textwidth}[t]{l@{\extracolsep{\fill}}r}
      \textbf{#1} & #2 \\
      \textit{\small#3} & \textit{\small #4} \\
    \end{tabular*}\vspace{-7pt}
}

\newcommand{\resumeSubSubheading}[2]{
    \item
    \begin{tabular*}{0.97\textwidth}{l@{\extracolsep{\fill}}r}
      \textit{\small#1} & \textit{\small #2} \\
    \end{tabular*}\vspace{-7pt}
}

\newcommand{\resumeProjectHeading}[2]{
    \item
    \begin{tabular*}{0.97\textwidth}{l@{\extracolsep{\fill}}r}
      \small#1 & #2 \\
    \end{tabular*}\vspace{-7pt}
}

\newcommand{\resumeSubItem}[1]{\resumeItem{#1}\vspace{-4pt}}

\renewcommand\labelitemii{$\vcenter{\hbox{\tiny$\bullet$}}$}

\newcommand{\resumeSubHeadingListStart}{\begin{itemize}[leftmargin=0.15in, label={}]}
\newcommand{\resumeSubHeadingListEnd}{\end{itemize}}
\newcommand{\resumeItemListStart}{\begin{itemize}}
\newcommand{\resumeItemListEnd}{\end{itemize}\vspace{-5pt}}

\begin{document}
"""

LATEX_END = r"\end{document}"


def _strip_urls(text: str) -> str:
    """Remove raw URLs - they contain & # ? which break LaTeX."""
    text = re.sub(r"https?://\S+", "", text)
    text = re.sub(r"www\.\S+", "", text)
    text = re.sub(r"  +", " ", text).strip()
    return text


def _escape_latex(text: str) -> str:
    if not text:
        return ""
    # Strip URLs first - they contain special chars that break LaTeX
    text = _strip_urls(text)
    replacements = [
        ("\\", r"\textbackslash{}"),
        ("&", r"\&"),
        ("%", r"\%"),
        ("$", r"\$"),
        ("#", r"\#"),
        ("^", r"\^{}"),
        ("_", r"\_"),
        ("{", r"\{"),
        ("}", r"\}"),
        ("~", r"\textasciitilde{}"),
    ]
    for old, new in replacements:
        text = text.replace(old, new)
    return text


def _build_header(name: str, contact: str) -> str:
    email = phone = linkedin = github = ""
    email_m = re.search(r"[\w.\-]+@[\w.\-]+", contact)
    if email_m:
        email = email_m.group()
    phone_m = re.search(r"[\+\(]?[\d\s\-\(\)]{7,}", contact)
    if phone_m:
        phone = phone_m.group().strip()
    linkedin_m = re.search(r"linkedin\.com/in/[\w\-]+", contact, re.I)
    if linkedin_m:
        linkedin = linkedin_m.group()
    github_m = re.search(r"github\.com/[\w\-]+", contact, re.I)
    if github_m:
        github = github_m.group()

    links = []
    if phone:
        links.append(_escape_latex(phone))
    if email:
        links.append(r"\href{mailto:" + email + r"}{" + _escape_latex(email) + r"}")
    if linkedin:
        links.append(
            r"\href{https://" + linkedin + r"}{" + _escape_latex(linkedin) + r"}"
        )
    if github:
        links.append(r"\href{https://" + github + r"}{" + _escape_latex(github) + r"}")

    link_line = r" $|$ ".join(links) if links else _escape_latex(contact)

    return (
        r"\begin{center}" + "\n"
        r"    \textbf{\Huge \scshape "
        + _escape_latex(name)
        + r"} \\ \vspace{1pt}"
        + "\n"
        r"    \small " + link_line + "\n"
        r"\end{center}" + "\n\n"
    )


def _build_summary(summary_lines: List[str]) -> str:
    if not summary_lines:
        return ""
    text = " ".join(l.strip() for l in summary_lines if l.strip())
    return (
        r"\section{Summary}" + "\n"
        r"\small{" + _escape_latex(text) + r"}" + "\n\n"
    )


def _build_experience(exp_lines: List[str]) -> str:
    if not exp_lines:
        return ""
    out = r"\section{Experience}" + "\n"
    out += r"  \resumeSubHeadingListStart" + "\n"
    jobs = _group_experience_lines(exp_lines)
    for job in jobs:
        company = _escape_latex(job.get("company", ""))
        title = _escape_latex(job.get("title", ""))
        date = _escape_latex(job.get("date", ""))
        location = _escape_latex(job.get("location", ""))
        bullets = job.get("bullets", [])
        out += (
            f"    \\resumeSubheading{{{company}}}{{{date}}}{{{title}}}{{{location}}}\n"
        )
        if bullets:
            out += "      \\resumeItemListStart\n"
            for b in bullets:
                b_clean = re.sub(r"^[•\-\*]\s*", "", b.strip())
                out += f"        \\resumeItem{{{_escape_latex(b_clean)}}}\n"
            out += "      \\resumeItemListEnd\n"
    out += r"  \resumeSubHeadingListEnd" + "\n\n"
    return out


def _group_experience_lines(lines: List[str]) -> List[Dict]:
    date_re = re.compile(
        r"\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|\d{4})\b.*"
        r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|\d{4}|Present|present|current|Current)",
        re.I,
    )
    jobs = []
    current_job: Dict = {}
    current_bullets: List[str] = []

    def save_job():
        if current_job:
            current_job["bullets"] = current_bullets[:]
            jobs.append(dict(current_job))

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        is_bullet = stripped.startswith(("•", "-", "*", "–"))
        if is_bullet:
            current_bullets.append(stripped)
        else:
            has_date = bool(date_re.search(stripped))
            has_pipe = "|" in stripped
            if has_date or has_pipe or (len(stripped) < 80 and not current_job):
                save_job()
                current_job = {}
                current_bullets = []
                if has_pipe:
                    parts = [p.strip() for p in stripped.split("|")]
                    current_job["company"] = parts[0] if len(parts) > 0 else stripped
                    current_job["title"] = parts[1] if len(parts) > 1 else ""
                    current_job["date"] = parts[2] if len(parts) > 2 else ""
                    current_job["location"] = parts[3] if len(parts) > 3 else ""
                else:
                    date_match = date_re.search(stripped)
                    if date_match:
                        before = (
                            stripped[: date_match.start()].strip().rstrip(",–-").strip()
                        )
                        current_job["date"] = date_match.group()
                        sub = re.split(r"\s{2,}|,\s*", before, 1)
                        current_job["company"] = sub[0].strip() if sub else before
                        current_job["title"] = sub[1].strip() if len(sub) > 1 else ""
                    else:
                        current_job["company"] = stripped
                        current_job["title"] = ""
                        current_job["date"] = ""
                    current_job["location"] = ""
            else:
                if current_job:
                    current_bullets.append(stripped)
                else:
                    current_job = {
                        "company": stripped,
                        "title": "",
                        "date": "",
                        "location": "",
                    }

    save_job()
    return jobs


def _build_education(edu_lines: List[str]) -> str:
    if not edu_lines:
        return ""
    out = r"\section{Education}" + "\n"
    out += r"  \resumeSubHeadingListStart" + "\n"
    i = 0
    while i < len(edu_lines):
        line = edu_lines[i].strip()
        if not line:
            i += 1
            continue
        date_str = ""
        m = re.search(r"(\d{4})\s*[-–]\s*(\d{4}|Present|present)", line)
        if m:
            date_str = m.group()
            line = line.replace(date_str, "").strip().rstrip(",–-").strip()
        school = line
        degree = ""
        if (
            i + 1 < len(edu_lines)
            and edu_lines[i + 1].strip()
            and not edu_lines[i + 1].strip().startswith(("•", "-", "*"))
        ):
            degree = edu_lines[i + 1].strip()
            i += 1
        out += (
            f"    \\resumeSubheading{{{_escape_latex(school)}}}{{{_escape_latex(date_str)}}}{{{_escape_latex(degree)}}}{{}}"
            + "\n"
        )
        i += 1
    out += r"  \resumeSubHeadingListEnd" + "\n\n"
    return out


def _build_skills(skill_lines: List[str]) -> str:
    if not skill_lines:
        return ""
    out = r"\section{Technical Skills}" + "\n"
    out += r" \begin{itemize}[leftmargin=0.15in, label={}]" + "\n"
    out += r"    \small{\item{" + "\n"
    categories = []
    for line in skill_lines:
        stripped = line.strip()
        if not stripped:
            continue
        if ":" in stripped:
            cat, _, items = stripped.partition(":")
            categories.append(
                r"     \textbf{"
                + _escape_latex(cat.strip())
                + r"}{: "
                + _escape_latex(items.strip())
                + r"} \\"
            )
        else:
            categories.append(
                r"     \textbf{Skills}{: " + _escape_latex(stripped) + r"} \\"
            )
    out += "\n".join(categories) + "\n"
    out += r"    }}" + "\n"
    out += r" \end{itemize}" + "\n\n"
    return out


def _build_projects(proj_lines: List[str]) -> str:
    if not proj_lines:
        return ""
    out = r"\section{Projects}" + "\n"
    out += r"    \resumeSubHeadingListStart" + "\n"
    parts = []
    current_name = ""
    current_bullets: List[str] = []

    def flush():
        nonlocal current_name, current_bullets
        if not current_name:
            return ""
        heading = r"\textbf{" + _escape_latex(current_name) + r"}"
        p = f"      \\resumeProjectHeading{{{heading}}}{{}}" + "\n"
        if current_bullets:
            p += "      \\resumeItemListStart\n"
            for b in current_bullets:
                b_clean = re.sub(r"^[•\-\*]\s*", "", b.strip())
                p += f"        \\resumeItem{{{_escape_latex(b_clean)}}}\n"
            p += "      \\resumeItemListEnd\n"
        return p

    for line in proj_lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith(("•", "-", "*")):
            current_bullets.append(stripped)
        else:
            parts.append(flush())
            current_name = stripped
            current_bullets = []
    parts.append(flush())

    out += "".join(p for p in parts if p)
    out += r"    \resumeSubHeadingListEnd" + "\n\n"
    return out


def _build_certifications(cert_lines: List[str]) -> str:
    if not cert_lines:
        return ""
    out = r"\section{Certifications}" + "\n"
    out += r" \begin{itemize}[leftmargin=0.15in, label={}]" + "\n"
    out += r"    \small{\item{" + "\n"
    items = [
        r"     \textbf{" + _escape_latex(l.strip().lstrip("•-* ")) + r"} \\"
        for l in cert_lines
        if l.strip()
    ]
    out += "\n".join(items) + "\n"
    out += r"    }}" + "\n"
    out += r" \end{itemize}" + "\n\n"
    return out


SECTION_KEYWORDS = {
    "summary": ["summary", "professional summary", "objective", "profile", "about"],
    "experience": [
        "experience",
        "work experience",
        "professional experience",
        "employment",
        "work history",
    ],
    "education": ["education", "academic", "degrees", "university", "college"],
    "skills": [
        "skills",
        "technical skills",
        "core competencies",
        "expertise",
        "technologies",
        "tools",
    ],
    "certifications": ["certifications", "certificates", "licenses", "accreditations"],
    "projects": ["projects", "personal projects", "side projects", "portfolio"],
}


def _parse_resume_sections(resume_text: str) -> Dict[str, Any]:
    sections: Dict[str, Any] = {"name": "", "contact": ""}
    lines = resume_text.split("\n")
    current_section = None
    current_content: List[str] = []

    def save_section():
        if current_section and current_content:
            sections[current_section] = current_content[:]

    for line in lines[:5]:
        stripped = line.strip()
        if stripped and len(stripped.split()) <= 5:
            if not any(
                x in stripped.lower()
                for x in ["@", "http", "phone", "email", "linkedin"]
            ):
                sections["name"] = stripped
                break

    contact_parts = []
    for line in lines[:10]:
        if any(
            x in line.lower() for x in ["@", "linkedin", "github", "phone", "http", "+"]
        ):
            contact_parts.append(line.strip())
    sections["contact"] = " | ".join(contact_parts)

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        clean_line = re.sub(r"[:\-\*=_#\|]+$", "", stripped.lower()).strip()
        matched_section = None
        for section_name, keywords in SECTION_KEYWORDS.items():
            if any(clean_line == kw or clean_line.startswith(kw) for kw in keywords):
                matched_section = section_name
                break
        if matched_section:
            save_section()
            current_section = matched_section
            current_content = []
        elif current_section:
            current_content.append(stripped)

    save_section()
    return sections


class ResumeGenerator:
    """Generate tailored resumes in DOCX and PDF (Jake's Resume style)"""

    def __init__(self):
        self.output_dir = OUTPUT_DIR
        os.makedirs(self.output_dir, exist_ok=True)

    def generate_pdf(
        self, resume_text: str, job_data: Dict[str, str], filename: str
    ) -> str:
        sections = _parse_resume_sections(resume_text)
        latex = LATEX_PREAMBLE
        latex += _build_header(sections.get("name", ""), sections.get("contact", ""))
        latex += _build_summary(sections.get("summary", []))
        latex += _build_experience(sections.get("experience", []))
        latex += _build_education(sections.get("education", []))
        latex += _build_skills(sections.get("skills", []))
        latex += _build_projects(sections.get("projects", []))
        latex += _build_certifications(sections.get("certifications", []))
        latex += LATEX_END
        output_path = os.path.join(self.output_dir, filename)
        self._compile_latex(latex, output_path)
        return output_path

    def _compile_latex(self, latex_source: str, output_pdf_path: str):
        with tempfile.TemporaryDirectory() as tmpdir:
            tex_file = os.path.join(tmpdir, "resume.tex")
            with open(tex_file, "w", encoding="utf-8") as f:
                f.write(latex_source)
            result = subprocess.run(
                [
                    "xelatex",
                    "-interaction=nonstopmode",
                    "-output-directory",
                    tmpdir,
                    tex_file,
                ],
                capture_output=True,
                text=True,
            )
            pdf_tmp = os.path.join(tmpdir, "resume.pdf")
            if os.path.exists(pdf_tmp):
                import shutil

                shutil.copy(pdf_tmp, output_pdf_path)
            else:
                tex_out = output_pdf_path.replace(".pdf", ".tex")
                with open(tex_out, "w", encoding="utf-8") as f:
                    f.write(latex_source)
                raise RuntimeError(
                    f"xelatex failed. LaTeX saved to {tex_out}.\n"
                    f"STDOUT: {result.stdout[-2000:]}"
                )

    def generate_docx(
        self, resume_text: str, job_data: Dict[str, str], filename: str
    ) -> str:
        sections = _parse_resume_sections(resume_text)
        doc = Document()

        for sec in doc.sections:
            sec.top_margin = Inches(0.5)
            sec.bottom_margin = Inches(0.5)
            sec.left_margin = Inches(0.5)
            sec.right_margin = Inches(0.5)

        self._set_default_style(doc)

        # Name
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(sections.get("name", ""))
        r.bold = True
        r.font.size = Pt(20)

        # Contact
        if sections.get("contact"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(sections["contact"]).font.size = Pt(10)

        # Summary
        if sections.get("summary"):
            self._add_section_header(doc, "Summary")
            text = " ".join(l.strip() for l in sections["summary"] if l.strip())
            doc.add_paragraph(text).paragraph_format.space_after = Pt(4)

        # Experience
        if sections.get("experience"):
            self._add_section_header(doc, "Experience")
            for job in _group_experience_lines(sections["experience"]):
                self._add_job_entry(doc, job)

        # Education
        if sections.get("education"):
            self._add_section_header(doc, "Education")
            edu = sections["education"]
            i = 0
            while i < len(edu):
                line = edu[i].strip()
                if not line:
                    i += 1
                    continue
                m = re.search(r"(\d{4})\s*[-–]\s*(\d{4}|Present)", line, re.I)
                date_str = m.group() if m else ""
                school = line.replace(date_str, "").strip().rstrip(",")
                degree = ""
                if (
                    i + 1 < len(edu)
                    and edu[i + 1].strip()
                    and not edu[i + 1].startswith(("•", "-", "*"))
                ):
                    degree = edu[i + 1].strip()
                    i += 1
                self._add_two_col_line(doc, school, date_str, bold_left=True)
                if degree:
                    self._add_two_col_line(doc, degree, "", italic_left=True)
                i += 1

        # Skills
        if sections.get("skills"):
            self._add_section_header(doc, "Technical Skills")
            for line in sections["skills"]:
                s = line.strip()
                if not s:
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                if ":" in s:
                    cat, _, items = s.partition(":")
                    r = p.add_run(cat.strip() + ": ")
                    r.bold = True
                    r.font.size = Pt(10)
                    p.add_run(items.strip()).font.size = Pt(10)
                else:
                    p.add_run(s).font.size = Pt(10)

        # Projects
        if sections.get("projects"):
            self._add_section_header(doc, "Projects")
            for line in sections["projects"]:
                s = line.strip()
                if not s:
                    continue
                if s.startswith(("•", "-", "*")):
                    bp = doc.add_paragraph(style="List Bullet")
                    bp.paragraph_format.left_indent = Inches(0.25)
                    bp.paragraph_format.space_after = Pt(2)
                    bp.add_run(s.lstrip("•-* ")).font.size = Pt(10)
                else:
                    p = doc.add_paragraph()
                    r = p.add_run(s)
                    r.bold = True
                    r.font.size = Pt(10)

        # Certifications
        if sections.get("certifications"):
            self._add_section_header(doc, "Certifications")
            for line in sections["certifications"]:
                s = line.strip().lstrip("•-* ")
                if s:
                    p = doc.add_paragraph(s)
                    p.paragraph_format.space_after = Pt(2)

        output_path = os.path.join(self.output_dir, filename)
        doc.save(output_path)
        return output_path

    def _set_default_style(self, doc):
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

    def _add_section_header(self, doc, title: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title.upper())
        r.bold = True
        r.font.size = Pt(12)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "000000")
        pBdr.append(bot)
        pPr.append(pBdr)

    def _add_two_col_line(
        self, doc, left: str, right: str, bold_left=False, italic_left=False
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), "9360")
        tabs.append(tab)
        pPr.append(tabs)
        r = p.add_run(left)
        r.bold = bold_left
        r.italic = italic_left
        r.font.size = Pt(10)
        if right:
            p.add_run("\t")
            rr = p.add_run(right)
            rr.italic = True
            rr.font.size = Pt(10)

    def _add_job_entry(self, doc, job: Dict):
        self._add_two_col_line(
            doc, job.get("company", ""), job.get("date", ""), bold_left=True
        )
        title = job.get("title", "")
        loc = job.get("location", "")
        if title or loc:
            self._add_two_col_line(doc, title, loc, italic_left=True)
        for bullet in job.get("bullets", []):
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.left_indent = Inches(0.25)
            bp.paragraph_format.space_after = Pt(2)
            bp.add_run(re.sub(r"^[•\-\*]\s*", "", bullet.strip())).font.size = Pt(10)

    def generate_both_formats(
        self, resume_text: str, job_data: Dict[str, str], base_filename: str
    ) -> Dict[str, str]:
        safe = re.sub(r"[/\\]", "_", base_filename)
        docx_filename = f"{safe}.docx"
        pdf_filename = f"{safe}.pdf"
        docx_path = self.generate_docx(resume_text, job_data, docx_filename)
        pdf_path = self.generate_pdf(resume_text, job_data, pdf_filename)
        return {
            "docx": docx_path,
            "pdf": pdf_path,
            "docx_name": docx_filename,
            "pdf_name": pdf_filename,
        }
