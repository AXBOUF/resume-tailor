import PyPDF2
from docx import Document
import re
from typing import Dict, List, Any

class ResumeParser:
    """Parse PDF and DOCX resumes to extract text and structure"""
    
    def __init__(self):
        self.sections = {
            'contact': [],
            'summary': [],
            'experience': [],
            'education': [],
            'skills': [],
            'certifications': [],
            'projects': [],
            'other': []
        }
    
    def parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            raise Exception(f"Error parsing PDF: {str(e)}")
        
        return self._structure_resume(text)
    
    def parse_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
        except Exception as e:
            raise Exception(f"Error parsing DOCX: {str(e)}")
        
        return self._structure_resume(text)
    
    def parse_file(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """Parse file based on type"""
        if file_type == 'pdf':
            return self.parse_pdf(file_path)
        elif file_type in ['docx', 'doc']:
            return self.parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def _structure_resume(self, text: str) -> Dict[str, Any]:
        """Attempt to structure the resume into sections"""
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        structured = {
            'raw_text': text,
            'lines': lines,
            'contact_info': self._extract_contact_info(lines),
            'sections': self._identify_sections(lines)
        }
        
        return structured
    
    def _extract_contact_info(self, lines: List[str]) -> Dict[str, str]:
        """Extract contact information from first few lines"""
        contact_info = {
            'name': '',
            'email': '',
            'phone': '',
            'location': '',
            'linkedin': '',
            'portfolio': ''
        }
        
        # Check first 10 lines for contact info
        for line in lines[:10]:
            # Email pattern
            email_match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', line)
            if email_match:
                contact_info['email'] = email_match.group(0)
            
            # Phone pattern
            phone_match = re.search(r'[\+]?[(]?[0-9]{3}[)]?[-\s\.]?[0-9]{3}[-\s\.]?[0-9]{4}', line)
            if phone_match:
                contact_info['phone'] = phone_match.group(0)
            
            # LinkedIn
            if 'linkedin' in line.lower():
                contact_info['linkedin'] = line
            
            # Portfolio/GitHub
            if any(x in line.lower() for x in ['github', 'portfolio', 'website']):
                contact_info['portfolio'] = line
        
        # Assume first non-empty line is name if it doesn't contain common non-name patterns
        for line in lines[:5]:
            if not any(x in line.lower() for x in ['@', 'http', 'phone', 'tel', 'email', 'linkedin']):
                if len(line.split()) <= 4 and not line.isdigit():
                    contact_info['name'] = line
                    break
        
        return contact_info
    
    def _identify_sections(self, lines: List[str]) -> Dict[str, List[str]]:
        """Identify common resume sections"""
        sections = {
            'summary': [],
            'experience': [],
            'education': [],
            'skills': [],
            'certifications': [],
            'projects': []
        }
        
        section_keywords = {
            'summary': ['summary', 'objective', 'profile', 'about'],
            'experience': ['experience', 'work experience', 'employment', 'professional experience', 'work history'],
            'education': ['education', 'academic', 'degree', 'university', 'college'],
            'skills': ['skills', 'technical skills', 'core competencies', 'expertise', 'technologies'],
            'certifications': ['certifications', 'certificates', 'licenses', 'accreditations'],
            'projects': ['projects', 'personal projects', 'side projects', 'portfolio']
        }
        
        current_section = None
        current_content = []
        
        for line in lines:
            # Check if line indicates a new section
            line_lower = line.lower()
            is_header = False
            
            for section, keywords in section_keywords.items():
                if any(keyword in line_lower for keyword in keywords):
                    # Save previous section if exists
                    if current_section and current_content:
                        sections[current_section] = current_content
                    
                    current_section = section
                    current_content = []
                    is_header = True
                    break
            
            if not is_header and current_section:
                current_content.append(line)
        
        # Save last section
        if current_section and current_content:
            sections[current_section] = current_content
        
        return sections